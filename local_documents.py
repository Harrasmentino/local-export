"""Local Word documents: compact reference guide, no external resources."""
import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Twips


def clean(value):
    return re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f]', '', str(value or ''))


def document(title):
    doc = Document()
    section = doc.sections[0]
    section.page_width, section.page_height = Inches(8.5), Inches(11)
    section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = Inches(1)
    section.header_distance = section.footer_distance = Inches(.492)
    for name, size, before, after, color in (
        ('Normal', 11, 0, 6, '202020'), ('Title', 24, 0, 10, '213F56'),
        ('Heading 1', 16, 18, 10, '2E74B5'), ('Heading 2', 13, 14, 7, '2E74B5'),
        ('Heading 3', 12, 10, 5, '1F4D78')):
        style = doc.styles[name]
        style.font.name, style.font.size = 'Calibri', Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        p = style.paragraph_format
        p.space_before, p.space_after, p.line_spacing = Pt(before), Pt(after), 1.25
        p.widow_control = True
    doc.styles['Normal'].paragraph_format.keep_with_next = False
    for name in ('Heading 1', 'Heading 2', 'Heading 3'):
        doc.styles[name].paragraph_format.keep_with_next = True
    section.header.paragraphs[0].text = 'Локальная выгрузка из Confluence'
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    field = OxmlElement('w:fldSimple')
    field.set(qn('w:instr'), 'PAGE')
    footer._p.append(field)
    doc.core_properties.author = ''
    doc.core_properties.last_modified_by = ''
    doc.add_paragraph(title, 'Title')
    doc.add_paragraph('Данные загружаются заново при каждом запуске. '
                      'Только Confluence; открытые источники и нормы ФАП автоматически не проверяются. '
                      'Неоднозначные правила требуют проверки человеком.')
    return doc


def table(doc, headers, rows, widths):
    result = doc.add_table(rows=1, cols=len(headers))
    result.autofit = False
    result.style = 'Table Grid'
    props = result._tbl.tblPr
    props.find(qn('w:tblW')).set(qn('w:type'), 'dxa')
    props.find(qn('w:tblW')).set(qn('w:w'), str(sum(widths)))
    indent = OxmlElement('w:tblInd')
    indent.set(qn('w:w'), '120')
    indent.set(qn('w:type'), 'dxa')
    props.append(indent)
    margins = OxmlElement('w:tblCellMar')
    for side, value in (('top', 80), ('bottom', 80), ('start', 120), ('end', 120)):
        item = OxmlElement('w:' + side)
        item.set(qn('w:w'), str(value))
        item.set(qn('w:type'), 'dxa')
        margins.append(item)
    props.append(margins)
    for col, width in zip(result.columns, widths):
        col.width = Twips(width)
    for row, values in zip([result.rows[0]] + [result.add_row() for _ in rows], [headers] + rows):
        for cell, value, width in zip(row.cells, values, widths):
            cell.width = Twips(width)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell.text = clean(value)
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(4)
                p.paragraph_format.line_spacing = 1.08
                for run in p.runs:
                    run.font.size = Pt(10)
    header_props = result.rows[0]._tr.get_or_add_trPr()
    header_props.append(OxmlElement('w:tblHeader'))
    for cell in result.rows[0].cells:
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), 'E8EEF5')
        cell._tc.get_or_add_tcPr().append(shading)
        for run in cell.paragraphs[0].runs:
            run.bold = True
    return result


def write_documents(records: list[dict], folder: Path):
    short = document('Неявка: таблица авиакомпаний')
    short.add_paragraph('Подробности и ссылки: neyavka_spravochnik.docx. Полные страницы и время загрузки: source_snapshot.json.')
    rows = [[r['airline'], r['iata'], r['moment'], r['status']] for r in records]
    table(short, ['Авиакомпания', 'ИАТА', 'Начало неявки', 'Статус проверки'], rows, [2400, 720, 2520, 3720])
    short.save(folder / 'neyavka_tablica.docx')

    guide = document('Неявка: справочник по источникам')
    guide.add_heading('Как читать', level=1)
    guide.add_paragraph('«Не найдено» означает, что автоматический поиск не выделил правило, а не что правила нет. '
                        'Число минут выводится только из простой однозначной формулировки целиком. '
                        'Прочие условия сохраняются текстом без автоматического толкования.')
    guide.add_paragraph('Источник и время загрузки каждой страницы сохранены в source_snapshot.json. '
                        'Тексты ниже — извлечённые фрагменты; для применения проверьте всю страницу.')
    for r in records:
        guide.add_heading(clean(r['airline']), level=1)
        guide.add_paragraph(clean(f'ИАТА: {r["iata"] or "не выделен"} | {r["moment"]}'))
        guide.add_paragraph(clean(r['status']))
        e = r['evidence']['noshow']
        guide.add_paragraph('Источник фрагмента: ' + e['origin'])
        text = e['text'] or 'Тематический фрагмент не найден. Откройте исходную статью.'
        for paragraph in re.split(r'\n\s*\n', text):
            guide.add_paragraph(clean(paragraph))
        source = guide.add_paragraph('Статья: ' + clean(r['url']))
        source.paragraph_format.space_before = Pt(4)
        source.paragraph_format.space_after = Pt(4)
    guide.save(folder / 'neyavka_spravochnik.docx')
